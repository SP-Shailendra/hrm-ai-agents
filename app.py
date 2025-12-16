import streamlit as st
import pandas as pd
import io
from docx import Document
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Image
import matplotlib.pyplot as plt
import re
from hrm_agent import (
    hr_feedback_agent,
    topic_modeling_agent,
    retention_report_agent,
    hr_chat_agent,
    hr_document_generator,
    create_pdf,
    performance_review_agent,
)

# Report Download Helpers

def generate_docx(report_text: str, figures: list | None = None):
    clean_text = clean_markdown(report_text)
    doc = Document()

    # Add text
    for line in clean_text.split("\n"):
        doc.add_paragraph(line)

    # Add figures
    if figures:
        doc.add_heading("Performance Visual Insights", level=2)
        for fig in figures:
            # fig here is a matplotlib figure, convert to buffer
            img_buf = fig_to_docx_image(fig)
            doc.add_picture(img_buf, width=Inches(4))
            doc.add_paragraph("")  # spacing

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_pdf(report_text: str, figures: list):
    clean_text = clean_markdown(report_text)
    buffer = io.BytesIO()
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(buffer)
    content = []

    # Add text
    for line in clean_text.split("\n"):
        content.append(Paragraph(line.replace("&", "&amp;"), styles["Normal"]))

    if figures:
        content.append(
            Paragraph("<br/><b>Performance Visual Insights</b><br/>", styles["Heading2"])
        )
        for fig in figures:
            content.append(fig)

    doc.build(content)
    buffer.seek(0)
    return buffer


def clean_markdown(text: str) -> str:
    text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text.strip()

def analyze_feedback(feedback_points: list[str]):
    strengths_keywords = [
        "ownership", "delivers", "adapts", "collaborates",
        "problem-solving", "scalable", "architecture", "structured"
    ]
    improvement_keywords = [
        "improve", "proactively", "communication", "updates"
    ]

    strengths = 0
    improvements = 0

    for fb in feedback_points:
        fb_lower = fb.lower()
        if any(k in fb_lower for k in strengths_keywords):
            strengths += 1
        if any(k in fb_lower for k in improvement_keywords):
            improvements += 1

    return strengths, max(improvements, 1)

def fig_to_image(fig):
    """
    Convert matplotlib figure to ReportLab Image
    """
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return Image(buf, width=350, height=250)  # SMALL & CLEAN

def fig_to_docx_image(fig):
    """
    Convert matplotlib figure to image buffer for DOCX
    """
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf

# ----------------------------------------------------
# PAGE CONFIG
# ----------------------------------------------------
st.set_page_config(page_title="HRM AI Agent", layout="wide")

st.title("👨‍💼 HRM Employee Feedback AI Agent")
st.write("Analyze employee feedback using AI — Sentiment, Topics, Retention & HR Chatbot")

# ----------------------------------------------------
# TABS
# ----------------------------------------------------
tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs(
    [
        "🧠 Sentiment Analysis",
        "📌 Topic Modeling",
        "📊 Retention Report",
        "💬 HR Chat Assistant",
        "📄 HR Document Generator",
        "📈 Performance & Goals",
        "🎧 HR Service Desk"
    ]
)

# ----------------------------------------------------
# TAB 1 - SENTIMENT ANALYSIS
# ----------------------------------------------------
with tab1:
    st.header("🧠 Employee Sentiment Analysis")
    text = st.text_area("Enter employee feedback:")
    if st.button("Analyze"):
        if text.strip():
            result = hr_feedback_agent(text)
            st.success(result)
        else:
            st.warning("Please enter feedback text!")

# ----------------------------------------------------
# TAB 2 - TOPIC MODELING
# ----------------------------------------------------
with tab2:
    st.header("📌 Topic Modeling on Multiple Feedback")
    feedbacks = st.text_area("Enter multiple feedback (one per line):")
    if st.button("Extract Topics"):
        lines = [f.strip() for f in feedbacks.split("\n") if f.strip()]
        if lines:
            result = topic_modeling_agent(lines)
            st.info(result)
        else:
            st.warning("Enter at least one feedback")

# ----------------------------------------------------
# TAB 3 - RETENTION REPORT
# ----------------------------------------------------
with tab3:
    st.header("📊 Employee Retention Report Generator")
    feedbacks_ret = st.text_area("Enter feedback list (one per line):")
    if st.button("Generate Report"):
        lines = [f.strip() for f in feedbacks_ret.split("\n") if f.strip()]
        if lines:
            result = retention_report_agent(lines)
            st.success(result)
        else:
            st.warning("Enter feedback")

# ----------------------------------------------------
# TAB 4 - HR CHATBOT
# ----------------------------------------------------
with tab4:
    st.header("💬 HR Service Assistant")
    query = st.text_input("Ask anything related to HR policies, employee satisfaction, etc.")
    if st.button("Ask"):
        if query.strip():
            answer = hr_chat_agent(query)
            st.write(answer)
        else:
            st.warning("Please type a question!")

# ----------------------------------------------------
# TAB 5 - HR DOCUMENT GENERATOR (FIXED)
# ----------------------------------------------------
with tab5:
    st.header("📄 HR Document Generator")

    # Company branding
    st.image("https://profisee.com/wp-content/uploads/2020/04/Logo_CompunnelDigital.png", width=160)
    st.markdown(
        """
        **Compunnel Digital**  
        📍 Noida, Uttar Pradesh, India  

        Compunnel Digital is a global digital transformation and IT services company
        helping organizations modernize technology and accelerate digital growth.
        """
    )
    
    st.divider()

    doc_type = st.selectbox(
        "Select Document Type",
        [
            "Offer Letter",
            "Appointment Letter",
            "Promotion Letter",
            "Salary Revision Letter",
            "Warning Letter"
        ]
    )

    employee_name = st.text_input("Employee Name")
    role = st.text_input("Role / Position")

    salary = None
    reason = None

    if doc_type in [
        "Offer Letter",
        "Appointment Letter",
        "Promotion Letter",
        "Salary Revision Letter"
    ]:
        salary = st.text_input("Salary Offered (Optional)")

    if doc_type == "Warning Letter":
        reason = st.text_area("Reason for Warning")

    if st.button("Generate Document"):
        if employee_name and role:
            # Generate document using existing function
            document_text = hr_document_generator(
                doc_type,
                employee_name,
                role,
                salary,
                reason
            )

            # Add company header manually (SAFE FIX)
            document_text = f"""
COMPANY DETAILS
Company Name: Compunnel Digital
Location: Noida, Uttar Pradesh, India

{document_text}
"""

            st.success("✅ Document Generated Successfully")
            st.text_area("Generated Document", document_text, height=350)

            pdf_file = create_pdf(
                document_text,
                filename=f"{doc_type.replace(' ', '_')}.pdf"
            )

            with open(pdf_file, "rb") as f:
                st.download_button(
                    "📥 Download PDF",
                    data=f,
                    file_name=f"{doc_type.replace(' ', '_')}.pdf",
                    mime="application/pdf"
                )
        else:
            st.warning("Please fill employee name and role!")

# ----------------------------------------------------
# TAB 6 - PERFORMANCE & GOALS
# ----------------------------------------------------
with tab6:
    st.header("📈 Performance & Goals Agent")

    emp_name = st.text_input("Employee Name", key="perf_name")
    emp_role = st.text_input("Employee Role", key="perf_role")
    kpis = st.text_area("KPIs / Expected Outcomes")
    emp_feedback = st.text_area(
        "Enter Performance Feedback (Manager + Peer Feedback)",
        height=200,
        key="perf_feedback"
    )

    # Initialize session state
    if "performance_result" not in st.session_state:
        st.session_state.performance_result = None

    if st.button("Generate Performance Review"):
        if emp_name and emp_role and emp_feedback:
            result = performance_review_agent(emp_name, emp_role, emp_feedback)
            st.session_state.performance_result = result
            st.success("✅ Performance Review Generated")
            st.write(result)
        else:
            st.warning("Please fill all fields.")

    # -------------------------------
    # DOWNLOAD SECTION
    # -------------------------------
    if st.session_state.performance_result:
        st.markdown("---")
        st.subheader("⬇️ Download Performance Review")

        col1, col2 = st.columns(2)

        with col1:
            docx_buffer = generate_docx(
                st.session_state.performance_result,
                figures=None  # plug figures later if needed
            )
            st.download_button(
                label="📄 Download as DOCX",
                data=docx_buffer,
                file_name=f"{emp_name}_Performance_Review.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with col2:
            pdf_buffer = generate_pdf(
                st.session_state.performance_result,
                figures=None
            )
            st.download_button(
                label="📑 Download as PDF",
                data=pdf_buffer,
                file_name=f"{emp_name}_Performance_Review.pdf",
                mime="application/pdf"
            )